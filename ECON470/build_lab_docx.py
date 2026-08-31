"""Build Lab 1 Problems 2 & 3 report as .docx (ECON 470)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parent


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    run.font.size = Pt(11)


def main():
    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = t.add_run("Lab 1: Retraining MicroGPT\nProblems 2 & 3 — Writeup")
    r.bold = True
    r.font.size = Pt(14)
    doc.add_paragraph("ECON 470 — Microeconomics of Artificial Intelligence")
    doc.add_paragraph("Program outputs are saved alongside this document:")
    doc.add_paragraph("• program_output_problem2_ces.txt", style="List Bullet")
    doc.add_paragraph("• program_output_microols.txt (readable training summary)", style="List Bullet")
    doc.add_paragraph("• program_output_microols_raw.txt (full raw stdout)", style="List Bullet")
    doc.add_paragraph()

    # ----- Problem 2 -----
    add_heading(doc, "Problem 2: Backpropagation and the Value class", 1)

    add_heading(doc, "4.1 Part A — What the Value class does", 2)
    p(doc, "(7) The Value class wraps each scalar in .data and records how it was computed: which other Values it depends on (children) and the partial derivative of this node with respect to each child (local gradients). Together, those links form a computation graph. Training needs ∂loss/∂(each parameter); storing children and local derivatives lets backward() apply the chain rule and fill in those derivatives automatically.")

    p(doc, "(8) Calling L.backward() topologically orders all nodes feeding into L, sets L.grad = 1, then walks backward. At each node it distributes gradient to children: each child.grad increases by (local ∂output/∂child) × (gradient w.r.t. this node). When finished, every parameter leaf has ∂L/∂parameter for the optimizer.")

    add_heading(doc, "4.2 Part B — CES production function (α = 0.5, ρ = 0.5, L = 4, K = 9)", 2)
    p(doc, "(9) Hand calculation for Y:", bold=True)
    p(doc, "Y = (0.5·L^ρ + (1−α)·K^ρ)^(1/ρ) with α = 0.5, ρ = 0.5.")
    p(doc, "Y = (0.5·4^0.5 + 0.5·9^0.5)^(1/0.5) = (0.5·2 + 0.5·3)^2 = 2.5^2 = 6.25.")

    p(doc, "(10) Let u = 0.5 L^0.5 + 0.5 K^0.5 and Y = u^2. Then dY/du = 2u, ∂u/∂L = 0.25 L^−0.5, ∂u/∂K = 0.25 K^−0.5.")
    p(doc, "At L = 4, K = 9: u = 2.5.")
    p(doc, "∂Y/∂L = 2u · 0.25 L^−0.5 = 2(2.5)(0.25)(1/2) = 0.625.")
    p(doc, "∂Y/∂K = 2u · 0.25 K^−0.5 = 2(2.5)(0.25)(1/3) = 5/12 ≈ 0.4167.")

    p(doc, "(11) Verification: run python3 problem2_ces.py. Output (see program_output_problem2_ces.txt): Y = 6.25, MP_L = 0.625, MP_K = 0.416666… — matches the hand calculations.")

    add_heading(doc, "4.3 Part C — Why MicroGPT needs this", 2)
    p(doc, "(12) Every weight must receive a gradient so the optimizer can change it to reduce loss. Wrapping weights and the loss in Value objects connects them in one graph from inputs to loss. loss.backward() computes ∂loss/∂(each parameter). The training loop uses those gradients (e.g. Adam) to update each p.data and improve predictions.")

    # ----- Problem 3 -----
    add_heading(doc, "Problem 3: MicroOLS — Regression by gradient descent", 1)

    add_heading(doc, "5.1 Part A — Read the code", 2)
    p(doc, "(13) Three blocks shared with microgpt.py: (1) the Value autodiff class and backward(); (2) Adam-style updates after backward(); (3) the overall train loop pattern (forward → loss → backward → optimizer).")

    p(doc, "(14) Three differences: (1) MicroOLS is a linear log–log demand model with brand dummies; MicroGPT is a transformer on tokens. (2) Loss is MSE on ln(sales) vs cross-entropy on next-token probabilities. (3) Inference samples continuous ln Q with Gaussian noise then exp, vs discrete softmax sampling over characters.")

    add_heading(doc, "5.2 Part B — Training output", 2)
    p(doc, "(15) From program_output_microols.txt (class microols.py, random.seed(42)):", bold=True)
    p(doc, "Step 1: loss 86.0754. Step 10000: loss 0.8775.")
    p(doc, "ln(Q) = −3.1207·ln(P) + 0.8523·MinuteMaid + 1.4959·Tropicana + 10.8151.")
    p(doc, "Demand elasticity a ≈ −3.12: |a| > 1 → demand is elastic (roughly ~3% quantity response to a 1% price change, holding brand). Negative a → higher price associated with lower sales.")
    p(doc, "b_mm and b_tr positive → at the same price, Minute Maid and especially Tropicana have higher predicted log sales than Dominick’s (omitted category).")
    p(doc, "Residual SE on ln(Q): 0.7938.")

    p(doc, "(16) Closed-form OLS from the lab: a = −3.1387, b_mm = 0.8702, b_tr = 1.5299, b0 = 10.8288. Gradient descent (Adam + minibatches) is within a few hundredths to tenths — close but not identical to exact OLS. The same autodiff + gradient machinery that scales to GPT also approximates the OLS solution for a linear model.")

    add_heading(doc, "5.3 Part C — Inference", 2)
    p(doc, "(17) At P = $3.00, Tropicana: predicted ln(Q) = 8.8825. Twenty sampled Q values (units) range from about 2,328 to 28,160 in this run. Wide range is expected: Gaussian noise on ln Q combined with exp produces multiplicative variation in levels.")

    p(doc, "(18) Analogous: both sample from a conditional distribution given inputs. Different: MicroGPT samples a discrete next token from softmax; MicroOLS adds Gaussian noise on log sales and exponentiates, giving a continuous skewed distribution for Q.")

    add_heading(doc, "5.4 Part D — Connecting to the course", 2)
    p(doc, "(19) MicroOLS and MicroGPT are both prediction technologies in the sense of Gans Ch. 1: they map inputs into a loss, minimize it with gradients, and yield fitted relationships or simulated draws. In MicroOLS, prediction is conditional mean (and implied distribution) of sales given price and brand; in MicroGPT, it is next-token probabilities given context. Judgment sits outside the optimizer: humans choose the target, functional form, loss, data, and how predictions inform decisions.")

    add_heading(doc, "Appendix — How outputs were captured", 1)
    p(doc, "Run: python3 capture_program_outputs.py (or run problem2_ces.py and microols.py manually). Raw microols stdout is large because training uses carriage return (\\r); program_output_microols.txt is a readable summary with first/last step and final sections.")

    out = ROOT / "Lab1_Problems_2_and_3_Report.docx"
    doc.save(out)
    print(f"Saved {out}")


if __name__ == "__main__":
    main()
